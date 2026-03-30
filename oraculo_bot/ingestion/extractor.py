"""Extrator de texto para o pipeline de ingestao RAG Lote 1.

Suporta extracao de texto de:
- PDF (texto nativo e OCR via pdfplumber/pytesseract)
- DOCX (via python-docx)
- TXT/MD (direto)
- RTF/ODT (fallback)

Especificacao: docs/RAG_LOTE1/IMPLEMENTATION_TODO.md (EPIC 3)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Optional


class MetodoExtracao(str, Enum):
    """Metodos de extracao de texto."""
    PDF_TEXTO = "pdf_texto"
    PDF_OCR = "pdf_ocr"
    PDF_HIBRIDO = "pdf_hibrido"
    DOCX = "docx"
    TEXTO = "texto"
    MARKDOWN = "markdown"
    RTF = "rtf"
    ODT = "odt"
    DESCONHECIDO = "desconhecido"


class StatusExtracao(str, Enum):
    """Status de extracao."""
    OK = "ok"
    PARCIAL = "parcial"
    FALHA = "falha"
    VAZIO = "vazio"
    CURTO_DEMAIS = "curto_demais"
    CORROMPIDO = "corrompido"
    QUARENTENA = "quarentena"


# Limites para deteccao de texto ruim
MIN_CHARS_TEXTO = 50  # Minimo de caracteres para texto util
MAX_RATIO_CHARS_QUEBRADOS = 0.3  # Maximo 30% de chars suspeitos

# Caracteres que indicam problema de codificacao ou OCR ruim
CHARS_QUEBRADOS = re.compile(r"[�\x00-\x08\x0b\x0c\x0e-\x1f]")


@dataclass
class ResultadoExtracao:
    """Resultado da extracao de texto de um arquivo."""
    caminho_arquivo: str
    texto: Optional[str] = None
    metodo: MetodoExtracao = MetodoExtracao.DESCONHECIDO
    status: StatusExtracao = StatusExtracao.OK
    char_count: int = 0
    pagina_count: Optional[int] = None
    erro: Optional[str] = None
    precisa_quarentena: bool = False
    motivo_quarentena: Optional[str] = None
    metadata: dict = field(default_factory=dict)
    extracted_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))


class TextoRuimDetector:
    """Detecta problemas em texto extraido."""

    @staticmethod
    def detectar_vazio(texto: Optional[str]) -> bool:
        """Verifica se o texto esta vazio ou apenas whitespace."""
        if not texto:
            return True
        return len(texto.strip()) == 0

    @staticmethod
    def detectar_curto_demais(texto: Optional[str], min_chars: int = MIN_CHARS_TEXTO) -> bool:
        """Verifica se o texto e muito curto."""
        if not texto:
            return True
        return len(texto.strip()) < min_chars

    @staticmethod
    def calcular_ratio_chars_quebrados(texto: str) -> float:
        """Calcula a proporcao de caracteres quebrados/problematicos."""
        if not texto:
            return 1.0

        total = len(texto)
        quebrados = len(CHARS_QUEBRADOS.findall(texto))

        return quebrados / total if total > 0 else 1.0

    @staticmethod
    def detectar_corrompido(texto: str, max_ratio: float = MAX_RATIO_CHARS_QUEBRADOS) -> bool:
        """Verifica se o texto parece corrompido."""
        ratio = TextoRuimDetector.calcular_ratio_chars_quebrados(texto)
        return ratio > max_ratio

    @staticmethod
    def avaliar_texto(texto: Optional[str]) -> tuple[StatusExtracao, bool, Optional[str]]:
        """Avalia qualidade do texto extraido.

        Args:
            texto: Texto extraido

        Returns:
            Tupla (status, precisa_quarentena, motivo)
        """
        if TextoRuimDetector.detectar_vazio(texto):
            return StatusExtracao.VAZIO, True, "Texto vazio"

        if TextoRuimDetector.detectar_curto_demais(texto):
            return StatusExtracao.CURTO_DEMAIS, True, f"Texto muito curto ({len(texto)} chars, min {MIN_CHARS_TEXTO})"

        if TextoRuimDetector.detectar_corrompido(texto):
            return StatusExtracao.CORROMPIDO, True, "Texto aparenta estar corrompido"

        return StatusExtracao.OK, False, None


class ExtratorTexto:
    """Extrator de texto de documentos com suporte a multiplos formatos."""

    def __init__(
        self,
        tentar_ocr: bool = True,
        min_chars_texto: int = MIN_CHARS_TEXTO,
        max_ratio_chars_quebrados: float = MAX_RATIO_CHARS_QUEBRADOS,
    ):
        """Inicializa o extrator.

        Args:
            tentar_ocr: Se deve tentar OCR quando PDF nao tiver texto nativo
            min_chars_texto: Minimo de caracteres para considerar texto valido
            max_ratio_chars_quebrados: Ratio maximo de chars quebrados aceitavel
        """
        self.tentar_ocr = tentar_ocr
        self.min_chars_texto = min_chars_texto
        self.max_ratio_chars_quebrados = max_ratio_chars_quebrados

    def extrair(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de um arquivo.

        Args:
            caminho: Caminho do arquivo

        Returns:
            ResultadoExtracao com texto e metadados
        """
        caminho_path = Path(caminho)

        if not caminho_path.exists():
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro="Arquivo nao existe",
                precisa_quarentena=True,
                motivo_quarentena="Arquivo nao encontrado",
            )

        extensao = caminho_path.suffix.lower()

        try:
            if extensao == ".pdf":
                return self._extrair_pdf(caminho)
            elif extensao == ".docx":
                return self._extrair_docx(caminho)
            elif extensao in {".txt", ".md"}:
                return self._extrair_texto(caminho)
            elif extensao == ".rtf":
                return self._extrair_rtf(caminho)
            elif extensao == ".odt":
                return self._extrair_odt(caminho)
            else:
                return ResultadoExtracao(
                    caminho_arquivo=caminho,
                    status=StatusExtracao.FALHA,
                    erro=f"Formato nao suportado: {extensao}",
                    precisa_quarentena=True,
                    motivo_quarentena=f"Extensao nao suportada: {extensao}",
                )

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro na extracao: {str(e)}",
            )

    def _avaliar_resultado(
        self,
        resultado: ResultadoExtracao,
    ) -> ResultadoExtracao:
        """Avalia qualidade do texto extraido e atualiza status.

        Args:
            resultado: Resultado da extracao bruta

        Returns:
            ResultadoExtracao com status atualizado
        """
        if resultado.status != StatusExtracao.OK:
            return resultado

        status, precisa_quarentena, motivo = TextoRuimDetector.avaliar_texto(resultado.texto)

        resultado.status = status
        resultado.char_count = len(resultado.texto) if resultado.texto else 0
        resultado.precisa_quarentena = precisa_quarentena
        resultado.motivo_quarentena = motivo

        return resultado

    def _extrair_pdf(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de PDF usando pdfplumber com fallback OCR.

        Args:
            caminho: Caminho do arquivo PDF

        Returns:
            ResultadoExtracao com texto extraido
        """
        try:
            import pdfplumber
        except ImportError:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro="pdfplumber nao instalado",
                precisa_quarentena=True,
                motivo_quarentena="Dependencia pdfplumber nao disponivel",
            )

        texto_paginas: list[str] = []
        total_paginas = 0
        paginas_com_texto = 0
        metodo = MetodoExtracao.PDF_TEXTO

        try:
            with pdfplumber.open(caminho) as pdf:
                total_paginas = len(pdf.pages)

                for pagina in pdf.pages:
                    texto = pagina.extract_text()
                    if texto and texto.strip():
                        texto_paginas.append(texto.strip())
                        paginas_com_texto += 1

            # Se nao conseguiu texto nativo, tentar OCR se habilitado
            if paginas_com_texto == 0 and self.tentar_ocr:
                resultado_ocr = self._extrair_pdf_ocr(caminho)
                if resultado_ocr.texto:
                    return self._avaliar_resultado(resultado_ocr)
                # Se OCR tambem falhou, retorna resultado vazio
                metodo = MetodoExtracao.PDF_OCR

            # Se algumas paginas tiveram texto mas nao todas, usar metodo hibrido
            elif paginas_com_texto < total_paginas and paginas_com_texto > 0:
                metodo = MetodoExtracao.PDF_HIBRIDO

            texto_final = "\n\n".join(texto_paginas)

            resultado = ResultadoExtracao(
                caminho_arquivo=caminho,
                texto=texto_final,
                metodo=metodo,
                pagina_count=total_paginas,
                metadata={
                    "paginas_total": total_paginas,
                    "paginas_com_texto": paginas_com_texto,
                },
            )

            return self._avaliar_resultado(resultado)

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.PDF_TEXTO,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro ao processar PDF: {str(e)}",
            )

    def _extrair_pdf_ocr(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de PDF usando OCR (pytesseract + pdf2image).

        Args:
            caminho: Caminho do arquivo PDF

        Returns:
            ResultadoExtracao com texto extraido via OCR
        """
        try:
            import pytesseract
            from pdf2image import convert_from_path
        except ImportError:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.PDF_OCR,
                erro="Dependencias OCR nao instaladas (pytesseract, pdf2image)",
            )

        try:
            imagens = convert_from_path(caminho, dpi=200)
            texto_paginas = []

            for imagem in imagens:
                texto = pytesseract.image_to_string(imagem, lang="por")
                if texto and texto.strip():
                    texto_paginas.append(texto.strip())

            texto_final = "\n\n".join(texto_paginas)

            return ResultadoExtracao(
                caminho_arquivo=caminho,
                texto=texto_final,
                metodo=MetodoExtracao.PDF_OCR,
                pagina_count=len(imagens),
                metadata={
                    "ocr_usado": True,
                    "dpi": 200,
                },
            )

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.PDF_OCR,
                erro=str(e),
            )

    def _extrair_docx(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de arquivo DOCX.

        Args:
            caminho: Caminho do arquivo DOCX

        Returns:
            ResultadoExtracao com texto extraido
        """
        try:
            from docx import Document
        except ImportError:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro="python-docx nao instalado",
                precisa_quarentena=True,
                motivo_quarentena="Dependencia python-docx nao disponivel",
            )

        try:
            doc = Document(caminho)
            paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            texto = "\n\n".join(paragrafos)

            # Tambem extrair texto de tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text.strip():
                            texto += f"\n{celula.text.strip()}"

            resultado = ResultadoExtracao(
                caminho_arquivo=caminho,
                texto=texto,
                metodo=MetodoExtracao.DOCX,
                metadata={
                    "paragrafos": len(doc.paragraphs),
                    "tabelas": len(doc.tables),
                },
            )

            return self._avaliar_resultado(resultado)

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.DOCX,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro ao processar DOCX: {str(e)}",
            )

    def _extrair_texto(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de arquivo TXT ou MD.

        Args:
            caminho: Caminho do arquivo de texto

        Returns:
            ResultadoExtracao com texto extraido
        """
        try:
            # Tentar diferentes encodings
            encodings = ["utf-8", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(caminho, "r", encoding=encoding) as f:
                        texto = f.read()

                    extensao = Path(caminho).suffix.lower()
                    metodo = MetodoExtracao.MARKDOWN if extensao == ".md" else MetodoExtracao.TEXTO

                    resultado = ResultadoExtracao(
                        caminho_arquivo=caminho,
                        texto=texto,
                        metodo=metodo,
                        metadata={"encoding": encoding},
                    )

                    return self._avaliar_resultado(resultado)

                except UnicodeDecodeError:
                    continue

            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.TEXTO,
                erro="Nao foi possivel decodificar arquivo",
                precisa_quarentena=True,
                motivo_quarentena="Encoding desconhecido",
            )

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.TEXTO,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro ao ler arquivo: {str(e)}",
            )

    def _extrair_rtf(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de arquivo RTF.

        Args:
            caminho: Caminho do arquivo RTF

        Returns:
            ResultadoExtracao com texto extraido
        """
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro="striprtf nao instalado",
                precisa_quarentena=True,
                motivo_quarentena="Dependencia striprtf nao disponivel",
            )

        try:
            with open(caminho, "r", encoding="utf-8") as f:
                rtf_content = f.read()

            texto = rtf_to_text(rtf_content)

            resultado = ResultadoExtracao(
                caminho_arquivo=caminho,
                texto=texto,
                metodo=MetodoExtracao.RTF,
            )

            return self._avaliar_resultado(resultado)

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.RTF,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro ao processar RTF: {str(e)}",
            )

    def _extrair_odt(self, caminho: str) -> ResultadoExtracao:
        """Extrai texto de arquivo ODT.

        Args:
            caminho: Caminho do arquivo ODT

        Returns:
            ResultadoExtracao com texto extraido
        """
        try:
            from odf.opendocument import load
            from odf import text
        except ImportError:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                erro="odfpy nao instalado",
                precisa_quarentena=True,
                motivo_quarentena="Dependencia odfpy nao disponivel",
            )

        try:
            doc = load(caminho)
            paragrafos = []

            for elemento in doc.getElementsByType(text.P):
                if elemento.firstChild:
                    paragrafos.append(str(elemento.firstChild))

            texto = "\n\n".join(paragrafos)

            resultado = ResultadoExtracao(
                caminho_arquivo=caminho,
                texto=texto,
                metodo=MetodoExtracao.ODT,
            )

            return self._avaliar_resultado(resultado)

        except Exception as e:
            return ResultadoExtracao(
                caminho_arquivo=caminho,
                status=StatusExtracao.FALHA,
                metodo=MetodoExtracao.ODT,
                erro=str(e),
                precisa_quarentena=True,
                motivo_quarentena=f"Erro ao processar ODT: {str(e)}",
            )
